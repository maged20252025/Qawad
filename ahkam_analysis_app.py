
import streamlit as st
import os
import docx
import fitz  # PyMuPDF لقراءة PDF
from docx import Document
import tempfile
import re

# ----- نموذج تبسيطي لاستخراج القواعد ----- #
def analyze_judgment(text):
    case_type = "مدنية" if "مدنية" in text else "غير محدد"
    issues = "لم تُستخرج تلقائيًا - النموذج بسيط"
    court_response = "لم يُحدد - النموذج التجريبي"
    legal_rule = "لم تُستخرج - يتطلب نموذج ذكي"
    return case_type, issues, court_response, legal_rule

# ----- قراءة PDF ----- #
def read_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

# ----- قراءة Word ----- #
def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

# ----- واجهة Streamlit ----- #
st.title("تحليل الأحكام القضائية واستخراج القواعد")
st.write("قم برفع ملفات الأحكام بصيغة PDF أو Word، وسنقوم بتحليلها تلقائيًا.")

uploaded_files = st.file_uploader("ارفع الأحكام هنا", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    if st.button("ابدأ التحليل وإنشاء الملف"):
        output_doc = Document()

        for uploaded_file in uploaded_files:
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name

            # استخراج رقم الطعن من اسم الملف
            file_name = uploaded_file.name
            match = re.findall(r"\d+", file_name)
            case_number = match[0] if match else "غير معروف"

            # استخراج النص من الملف
            if file_name.endswith(".pdf"):
                full_text = read_pdf(tmp_path)
            else:
                full_text = read_docx(tmp_path)

            # تحليل الحكم
            case_type, issues, court_response, legal_rule = analyze_judgment(full_text)

            # كتابة النتائج في ملف Word
            output_doc.add_paragraph(f"رقم الطعن: {case_number}")
            output_doc.add_paragraph(f"نوع القضية: {case_type}")
            output_doc.add_paragraph(f"الإشكاليات المثارة: {issues}")
            output_doc.add_paragraph(f"رد المحكمة العليا: {court_response}")
            output_doc.add_paragraph(f"القاعدة القضائية المستخلصة: {legal_rule}")
            output_doc.add_page_break()

            os.unlink(tmp_path)  # حذف الملف المؤقت

        # حفظ الملف النهائي
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
            output_doc.save(f.name)
            st.success("تم إنشاء الملف بنجاح!")
            with open(f.name, "rb") as file:
                st.download_button("📥 تحميل الملف النهائي", file.read(), file_name="القواعد_القضائية.docx")
